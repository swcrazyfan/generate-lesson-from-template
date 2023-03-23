import openai
import os
import streamlit as st
import csv
from datetime import datetime
import base64
from io import BytesIO
import docx
import streamlit.components.v1 as components

def log_to_csv(prompt, generated_content):
	with open("lesson_plan_logs.csv", mode="a", newline="", encoding="utf-8") as log_file:
		log_writer = csv.writer(log_file)
		log_writer.writerow([datetime.now(), prompt, generated_content])

def generate_content_from_template(user_prompt):
    	openai.api_key = os.environ["OPENAI_API_KEY"]

   	 full_prompt = (
        "Generate a well-structured lesson plan based on the user's input. Replace # with the correct enumeration.\n\n"
        f"Prompt: {user_prompt}\n\n"
        "Title (Total Estimated Time):\n"
        "#. Objective:\n"
        "#. Steps: Organize lesson plan activities under Steps as sub-sections. Each activity should contain estimated time, materials (if applicable), specific steps (using numbered list), and 1 reflection question. Be conservative with time estimates meaning overestimate or give a range.\n"
        "#.Closing:\n"
        "#. Materials: List any materials mentioned in the lesson plan.\n\n"
        "Ensure that the formatting and structure of the generated content are compatible with python-docx, and adhere to the standard Microsoft Word list format ('1., a., i.') for enumerating the sections, activities, and steps."
   	 )
	
	response = openai.ChatCompletion.create(
		model="gpt-3.5-turbo",
		messages=[
			{
				"role": "system",
				"content": (
					"You are a helpful assistant that can generate lesson plans based on a template and user prompt."
				),
			},
			{"role": "user", "content": full_prompt},
		],
		max_tokens=2000,
		n=1,
		stop=None,
		temperature=0.8,
	)

	generated_content = response["choices"][0]["message"]["content"].split("\n\n")

	# Replace abbreviations and placeholders with their full forms
	for i, item in enumerate(generated_content):
		for abbr, full in abbreviations.items():
			generated_content[i] = item.replace(abbr, full)
		for placeholder, replacement in placeholders.items():
			generated_content[i] = generated_content[i].replace(placeholder, replacement)

	return generated_content

def docx_from_generated_content(generated_content):
    document = docx.Document()

    for i, item in enumerate(generated_content):
        if i == 0:
            document.add_heading(item, level=1)
        else:
            heading_and_text = item.split("\n", 1)
            if len(heading_and_text) == 2:
                heading, text = heading_and_text
                document.add_heading(heading, level=2)
                document.add_paragraph(text.strip())
            elif len(heading_and_text) == 1 and heading_and_text[0].strip():
                # Only add non-empty strings as paragraphs
                document.add_paragraph(heading_and_text[0].strip())

    return document


st.title("Lesson Plan Generator")

user_prompt = st.text_input("Enter a prompt to guide the content generation:")

if st.button("Generate Lesson Plan"):
    with st.spinner("Generating..."):
        generated_content = generate_content_from_template(user_prompt)

        docx_document = docx_from_generated_content(generated_content)

        buffer = BytesIO()
        docx_document.save(buffer)
        buffer.seek(0)

        b64 = base64.b64encode(buffer.getvalue()).decode()

        lesson_name = "generated_lesson_plan"
        for line in generated_content:
            if line.startswith("Title:"):
                lesson_name = lesson_name.split("\n")[0]
                break

        st.download_button(
            label="Download Generated Lesson Plan",
            data=buffer,
            file_name=f"{lesson_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        log_to_csv(user_prompt, generated_content)
